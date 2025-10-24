"""Word document generator using python-docx."""

from pathlib import Path
from typing import Dict, Any
import re
from docx import Document
from .base import DocumentGenerator


class WordGenerator(DocumentGenerator):
    """Generator for Microsoft Word documents."""
    
    def generate(self, template_path: Path, fields: Dict[str, Any], output_path: Path) -> Path:
        """
        Generate a Word document from a template by replacing placeholders.
        
        Placeholders in the template should be in the format {{field_name}}.
        """
        # Load the template
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, fields)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, fields)
        
        # Replace placeholders in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_text_in_paragraph(paragraph, fields)
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_text_in_paragraph(paragraph, fields)
        
        # Save the document
        doc.save(output_path)
        return output_path
    
    def _replace_text_in_paragraph(self, paragraph, fields: Dict[str, Any]) -> None:
        """Replace placeholder text in a paragraph."""
        for field_name, field_value in fields.items():
            placeholder = f"{{{{{field_name}}}}}"
            if placeholder in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(field_value))
    
    def get_template_fields(self, template_path: Path) -> list[str]:
        """Extract field names from a Word template."""
        doc = Document(template_path)
        fields = set()
        
        # Find placeholders in format {{field_name}}
        pattern = r'\{\{([^}]+)\}\}'
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            fields.update(matches)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        fields.update(matches)
        
        # Search in headers and footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                fields.update(matches)
            for paragraph in section.footer.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                fields.update(matches)
        
        return sorted(list(fields))
